import os
import boto3
import streamlit as st
from docx import Document
from nltk.tokenize import sent_tokenize
import nltk
import tempfile

# Add new imports
try:
    import win32com.client
    import pythoncom
    WORD_AVAILABLE = True
except ImportError:
    WORD_AVAILABLE = False

# Ensure nltk data is available
try:
    nltk.data.find('tokenizers/punkt')
    nltk.data.find('tokenizers/punkt_tab')
except LookupError:
    nltk.download('punkt')
    nltk.download('punkt_tab')

class DocumentTranslator:
    def __init__(self, aws_key_id, aws_secret_key):
        self.translate_client = boto3.client(
            service_name='translate',
            region_name='us-east-1',
            aws_access_key_id=aws_key_id,
            aws_secret_access_key=aws_secret_key
        )

    def translate_text(self, text: str, source_lang: str, target_lang: str) -> str:
        if not text.strip():
            return text
            
        try:
            max_chunk_size = 5000
            sentences = sent_tokenize(text)
            chunks = []
            current_chunk = []
            current_size = 0

            for sentence in sentences:
                if current_size + len(sentence) > max_chunk_size:
                    chunks.append(' '.join(current_chunk))
                    current_chunk = []
                    current_size = 0
                current_chunk.append(sentence)
                current_size += len(sentence)

            if current_chunk:
                chunks.append(' '.join(current_chunk))

            translated_chunks = []
            for chunk in chunks:
                response = self.translate_client.translate_text(
                    Text=chunk,
                    SourceLanguageCode=source_lang,
                    TargetLanguageCode=target_lang
                )
                translated_chunks.append(response['TranslatedText'])

            return ' '.join(translated_chunks)
        except Exception as e:
            st.error(f"Translation error: {str(e)}")
            return text

    def translate_docx(self, input_docx_path: str, source_lang: str, target_lang: str, output_docx_path: str, progress_callback):
        doc = Document(input_docx_path)

        total_paragraphs = len(doc.paragraphs) + sum(len(table.rows) for table in doc.tables)
        progress_callback(0)  # Initialize progress bar

        # Process paragraphs
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                translated_text = self.translate_text(paragraph.text, source_lang, target_lang)
                for run in paragraph.runs:
                    run.text = translated_text
            progress_callback(int((i + 1) / total_paragraphs * 100))  # Update progress bar

        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        cell.text = self.translate_text(cell.text, source_lang, target_lang)
            progress_callback(100)  # 100% after processing tables

        doc.save(output_docx_path)

    def pdf_to_docx(self, pdf_path, docx_path, progress_callback):
        """Convert PDF to DOCX using Microsoft Word"""
        try:
            if WORD_AVAILABLE:
                # Initialize COM in the current thread
                pythoncom.CoInitialize()
                
                # Create Word application instance
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False

                try:
                    status_text = st.empty()
                    status_text.text("Opening PDF with Microsoft Word...")
                    
                    # Convert PDF to DOCX
                    doc = word.Documents.Open(pdf_path)
                    progress_callback(50)
                    
                    status_text.text("Saving as DOCX...")
                    doc.SaveAs2(docx_path, FileFormat=16)  # 16 = DOCX format
                    progress_callback(90)
                    
                    doc.Close()
                    progress_callback(100)
                    status_text.empty()
                    
                except Exception as e:
                    st.error(f"MS Word conversion error: {str(e)}")
                    raise
                finally:
                    word.Quit()
                    pythoncom.CoUninitialize()
            else:
                # Fallback to libreoffice
                import subprocess
                subprocess.run(['soffice', '--headless', '--convert-to', 'docx', pdf_path, '--outdir', os.path.dirname(docx_path)], check=True)
                progress_callback(100)

        except Exception as e:
            st.error(f"PDF conversion error: {str(e)}")
            raise

# Streamlit Application
st.set_page_config(page_title="DocuX Translator", layout="wide")
st.title("DocuX")

# Define more detailed language options
LANGUAGES = {
    "English": "en",
    "French": "fr",
    "Spanish": "es",
    "German": "de",
    "Italian": "it",
    "Arabic": "ar",
    "Chinese (Simplified)": "zh",
    "Japanese": "ja"
}

aws_key_id = os.environ.get('aws_key_id')
aws_secret_key = os.environ.get('aws_secret_key')

if not aws_key_id or not aws_secret_key:
    st.error("AWS credentials not found. Please set environment variables.")
else:
    translator = DocumentTranslator(aws_key_id, aws_secret_key)

    source_language = st.selectbox("Source Language", list(LANGUAGES.keys()), index=0)
    target_language = st.selectbox("Target Language", list(LANGUAGES.keys()), index=1)

    uploaded_file = st.file_uploader("Upload PDF or DOCX", type=["pdf", "docx"])

    if uploaded_file:
        # Create temporary directory for processing
        with tempfile.TemporaryDirectory() as temp_dir:
            try:
                input_path = os.path.join(temp_dir, uploaded_file.name)
                with open(input_path, "wb") as f:
                    f.write(uploaded_file.getvalue())

                progress_bar = st.progress(0)
                status_text = st.empty()

                if st.button("Translate"):
                    try:
                        status_text.text("Processing document...")
                        output_filename = f"translated_{uploaded_file.name}_{target_language}"
                        if not output_filename.endswith('.docx'):
                            output_filename = output_filename + '.docx'
                        output_path = os.path.join(temp_dir, output_filename)

                        source_lang_code = LANGUAGES[source_language]
                        target_lang_code = LANGUAGES[target_language]

                        if uploaded_file.type == "application/pdf":
                            temp_docx = os.path.join(temp_dir, "temp_converted.docx")
                            status_text.text("Converting PDF to DOCX...")
                            translator.pdf_to_docx(input_path, temp_docx, progress_bar.progress)
                            
                            status_text.text("Translating document...")
                            translator.translate_docx(temp_docx, source_lang_code, target_lang_code, output_path, progress_bar.progress)
                        else:
                            status_text.text("Translating document...")
                            translator.translate_docx(input_path, source_lang_code, target_lang_code, output_path, progress_bar.progress)

                        with open(output_path, "rb") as f:
                            st.download_button(
                                label="Download Translated Document",
                                data=f.read(),
                                file_name=output_filename,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )

                        status_text.text("Translation complete!")
                        progress_bar.progress(100)

                    except Exception as e:
                        st.error(f"Error during translation: {str(e)}")
                        progress_bar.empty()
                        status_text.empty()

            except Exception as e:
                st.error(f"Error processing file: {str(e)}")
